"""
把AI周报的Markdown文本转换成Word文档（.docx）。

这里不是一个"通用Markdown转Word"工具，而是专门针对我们自己
ai_report.py生成的报告格式写的（标题用##、表格用|分隔、加粗用**、
链接用[文字](网址)、各大段之间用 "\n\n---\n\n" 分隔），因为格式是
我们自己定的、完全可控，针对性解析比写一个通用转换器可靠、简单很多。

字体规则：英文/数字用Calibri，中文用仿宋——Word本身支持"中西文分体"，
只要底层XML的w:ascii（西文）和w:eastAsia（中文）设成不同字体，
同一段混排文字里，Word会自动按每个字符是不是中文来分别套用对应字体，
不需要自己在文本里手动拆分中英文。

标题/小标题额外加了蓝色（HEADING_COLOR），正文保持默认黑色。
"""
from __future__ import annotations

import re
from io import BytesIO

LATIN_FONT = "Calibri(西文标题)"      # 英文/数字用这个字体
CHINESE_FONT = "仿宋"        # 中文用这个字体
HEADING_COLOR = (0x1F, 0x4E, 0x79)  # 标题用的蓝色
TITLE_FONT_SIZE = 21
HEADING_FONT_SIZE = 15
BODY_FONT_SIZE = 11

def _add_bottom_border(paragraph, color_hex: str = "1F4E79", size: int = 12):
    """给段落底部加一条边框线，效果就是"这段文字下面一条横线"。
    size是线的粗细，单位是"1/8磅"，12大概是1.5磅粗，比较适合当标题下划线。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
 
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")   # 文字和线之间留一点空隙
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
 

def _set_font(run, size: int, bold: bool = False, color: tuple = None):
    """英文走w:ascii/w:hAnsi（Calibri），中文走w:eastAsia（仿宋）。
    color传(R,G,B)三元组就设置文字颜色，不传就是默认黑色。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    run.font.size = Pt(size)
    run.font.name = LATIN_FONT  # 这一步会顺带设置w:ascii和w:hAnsi
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CHINESE_FONT)


def _add_hyperlink(paragraph, url: str, text: str, size: int):
    """python-docx本身不支持直接加超链接，用官方文档认可的底层XML写法实现。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))  # Word的sz单位是"半磅"
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _render_rich_text(paragraph, text: str, size: int = BODY_FONT_SIZE):
    """把一段可能带**加粗**和[文字](url)链接的文本，写进某个paragraph里，
    链接是真的可以点击的超链接。同时用于普通段落和表格单元格。"""
    tokens = re.split(r"(\[[^\]]+\]\([^)]+\))", text)
    for token in tokens:
        if not token:
            continue
        link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link_match:
            link_text, url = link_match.groups()
            _add_hyperlink(paragraph, url, link_text, size)
            continue

        bold_parts = re.split(r"(\*\*[^*]+\*\*)", token)
        for part in bold_parts:
            if not part:
                continue
            is_bold = part.startswith("**") and part.endswith("**")
            run_text = part[2:-2] if is_bold else part
            run = paragraph.add_run(run_text)
            _set_font(run, size, bold=is_bold)


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line))


def markdown_report_to_docx(content: str, title: str) -> bytes:
    """把report的markdown文本转成docx，返回文件的二进制内容（可以直接用在
    Streamlit的st.download_button里）"""
    from docx import Document

    doc = Document()

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    _set_font(title_run, TITLE_FONT_SIZE, bold=True, color=HEADING_COLOR)
    _add_bottom_border(title_para)

    sections = content.split("\n\n---\n\n")
    for section in sections:
        lines = [l for l in section.strip().split("\n")]
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if line.startswith("## "):
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(line[3:].strip())
                _set_font(heading_run, HEADING_FONT_SIZE, bold=True, color=HEADING_COLOR)
                i += 1
                continue

            if line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                rows = [
                    [c.strip() for c in l.strip("|").split("|")]
                    for l in table_lines
                    if not _is_table_separator(l)
                ]
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    try:
                        table.style = "Light Grid Accent 1"
                    except KeyError:
                        pass
                    for r, row in enumerate(rows):
                        for c, cell_text in enumerate(row):
                            cell = table.cell(r, c)
                            cell.text = ""
                            _render_rich_text(cell.paragraphs[0], cell_text, size=BODY_FONT_SIZE - 1)
                continue

            para = doc.add_paragraph()
            _render_rich_text(para, line)
            if re.match(r"^\*\*[^*]+\*\*$", line):
                from docx.shared import Pt
                para.paragraph_format.space_after = Pt(2)
            i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# def _is_table_separator(line: str) -> bool:
#     """判断是不是markdown表格里的分隔线，比如 |---|---|---|"""
#     return bool(re.match(r"^\|?[\s\-:|]+\|?$", line))


# def _strip_markdown_link(text: str) -> str:
#     """把 [文字](网址) 转成 "文字 (网址)"，Word文本里不需要markdown语法"""
#     return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


# def _add_paragraph_with_bold(doc, text: str):
#     """普通段落，同时处理**加粗**和[链接](url)两种markdown语法"""
#     text = _strip_markdown_link(text)
#     para = doc.add_paragraph()
#     parts = re.split(r"(\*\*[^*]+\*\*)", text)
#     for part in parts:
#         if not part:
#             continue
#         if part.startswith("**") and part.endswith("**"):
#             run = para.add_run(part[2:-2])
#             run.bold = True
#         else:
#             para.add_run(part)


# def markdown_report_to_docx(content: str, title: str) -> bytes:
#     """把report的markdown文本转成docx，返回文件的二进制内容（可以直接用在
#     Streamlit的st.download_button里）"""
#     from docx import Document

#     doc = Document()
#     doc.add_heading(title, level=0)

#     sections = content.split("\n\n---\n\n")
#     for section in sections:
#         lines = [l for l in section.strip().split("\n")]
#         i = 0
#         while i < len(lines):
#             line = lines[i].strip()

#             if not line:
#                 i += 1
#                 continue

#             if line.startswith("## "):
#                 doc.add_heading(line[3:].strip(), level=1)
#                 i += 1
#                 continue

#             if line.startswith("|"):
#                 # 收集连续的表格行
#                 table_lines = []
#                 while i < len(lines) and lines[i].strip().startswith("|"):
#                     table_lines.append(lines[i].strip())
#                     i += 1
#                 rows = [
#                     [c.strip() for c in l.strip("|").split("|")]
#                     for l in table_lines
#                     if not _is_table_separator(l)
#                 ]
#                 if rows:
#                     table = doc.add_table(rows=len(rows), cols=len(rows[0]))
#                     try:
#                         table.style = "Light Grid Accent 1"
#                     except KeyError:
#                         pass  # 有些精简版Word模板没有这个样式，跳过样式，不影响内容
#                     for r, row in enumerate(rows):
#                         for c, cell_text in enumerate(row):
#                             table.cell(r, c).text = _strip_markdown_link(cell_text)
#                 continue

#             # 普通段落（可能带**加粗**、[链接](url)）
#             _add_paragraph_with_bold(doc, line)
#             i += 1

#     buf = BytesIO()
#     doc.save(buf)
#     return buf.getvalue()